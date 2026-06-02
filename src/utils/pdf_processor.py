"""
PDF/DOCX Document Processor
Handles resume and document text extraction.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes using pdfplumber (best quality)."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts).strip()
    except Exception as e:
        logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
        return _extract_pdf_fallback(file_bytes)


def _extract_pdf_fallback(file_bytes: bytes) -> str:
    """Fallback PDF extraction using PyPDF2."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        return "\n\n".join(
            page.extract_text() or "" for page in reader.pages
        ).strip()
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text(file_bytes: bytes, file_name: str) -> str:
    """Auto-detect file type and extract text."""
    name_lower = file_name.lower()
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name_lower.endswith(".txt"):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    else:
        logger.warning(f"Unsupported file type: {file_name}")
        return ""


def validate_file(file_bytes: bytes, file_name: str, max_size_mb: int = 10) -> tuple[bool, str]:
    """Validate file size and type."""
    allowed_extensions = {".pdf", ".docx", ".txt"}
    ext = "." + file_name.split(".")[-1].lower()
    if ext not in allowed_extensions:
        return False, f"File type '{ext}' not supported. Use PDF, DOCX, or TXT."
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return False, f"File too large ({size_mb:.1f}MB). Max is {max_size_mb}MB."
    return True, "OK"
